from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/superpowers/specs/2026-08-11-production-launch-design.md"
OUTPUT = ROOT / "docs/EZLove-项目立项及生产落地方案.docx"

FONT = "Arial Unicode MS"
MONO = "Arial Unicode MS"
INK = RGBColor(36, 48, 58)
BLUE = RGBColor(31, 78, 121)
TEAL = RGBColor(23, 107, 104)
MUTED = RGBColor(100, 112, 121)
LIGHT = "EAF2F5"
GRID = "C9D4DA"
USABLE_DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=10.5, bold=None, color=INK, mono=False):
    name = MONO if mono else FONT
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def add_inline(paragraph, text, size=10.5, color=INK):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            set_run_font(paragraph.add_run(part[2:-2]), size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size - 0.5, color=TEAL, mono=True)
            run.font.highlight_color = None
        else:
            link = re.fullmatch(r"\[(.*?)\]\((.*?)\)", part)
            if link:
                run = paragraph.add_run(link.group(1))
                set_run_font(run, size=size, color=BLUE)
                run.underline = True
            else:
                set_run_font(paragraph.add_run(part), size=size, color=color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, TEAL, 9, 4),
        ("Heading 4", 10.5, TEAL, 7, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(92)
    p.paragraph_format.space_after = Pt(14)
    set_run_font(p.add_run("EZLove 易挂念"), size=30, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    set_run_font(p.add_run("项目立项及生产落地方案"), size=20, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("让牵挂被看见"), size=13, color=MUTED)

    for label, value in (
        ("版本", "1.0"),
        ("日期", "2026-08-11"),
        ("部署边界", "阿里云中国内地，最多 1 台弹性云服务器"),
        ("首期范围", "溪东社区"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(f"{label}："), size=10.5, bold=True, color=MUTED)
        set_run_font(p.add_run(value), size=10.5, color=INK)
    doc.add_page_break()


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [USABLE_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for idx, line in enumerate(lines):
        run = p.add_run(line + ("\n" if idx < len(lines) - 1 else ""))
        set_run_font(run, size=8.3, color=INK, mono=True)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    normalized = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=cols)
    table.style = "Table Grid"
    if cols == 2:
        widths = [2300, 7060]
    elif cols == 3:
        widths = [1800, 2200, 5360]
    elif cols == 4:
        widths = [1450, 2200, 3600, 2110]
    elif cols == 5:
        widths = [1150, 1850, 2650, 2050, 1660]
    else:
        widths = [USABLE_DXA // cols] * cols
        widths[-1] += USABLE_DXA - sum(widths)
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(normalized):
        row_pr = table.rows[r_idx]._tr.get_or_add_trPr()
        row_pr.append(OxmlElement("w:cantSplit"))
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if c_idx == 1 and len(value) < 14:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, value, size=9.2, color=INK)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("start", "1"), ("numFmt", "decimal"), ("lvlText", "%1."), ("lvlJc", "left")):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:val"), value)
        level.append(node)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "548")
    ind.set(qn("w:hanging"), "274")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_paragraph(doc, text, num_id):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)
    p_pr.append(num_pr)
    add_inline(paragraph, text)


def parse_markdown(doc, text):
    lines = text.splitlines()
    idx = 0
    in_code = False
    code_lines = []
    current_num_id = None
    while idx < len(lines):
        raw = lines[idx]
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(line)
            idx += 1
            continue
        if not line.strip() or line.strip() == "---":
            current_num_id = None
            idx += 1
            continue
        if line.startswith("# "):
            idx += 1
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1)) - 1
            title = heading.group(2)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, title, size={1: 16, 2: 13, 3: 11.5}.get(level, 10.5), color=BLUE if level < 3 else TEAL)
            idx += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx].strip())
                idx += 1
            rows = []
            for t_line in table_lines:
                cells = [c.strip() for c in t_line.strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                    continue
                rows.append(cells)
            add_table(doc, rows)
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, line[2:], size=10, color=MUTED)
            idx += 1
            continue
        bullet = re.match(r"^\s*-\s+(.+)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            idx += 1
            continue
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if numbered:
            if current_num_id is None:
                current_num_id = create_numbering(doc)
            add_numbered_paragraph(doc, numbered.group(1), current_num_id)
            idx += 1
            continue
        current_num_id = None
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        add_inline(p, line)
        idx += 1
    if code_lines:
        add_code_block(doc, code_lines)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run("EZLove 易挂念｜项目立项及生产落地方案"), size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    add_cover(doc)
    source_text = SOURCE.read_text(encoding="utf-8")
    parse_markdown(doc, source_text[source_text.index("## A. 立项背景") :])

    core = doc.core_properties
    core.title = "EZLove 易挂念项目立项及生产落地方案"
    core.subject = "溪东社区政府项目立项与生产落地"
    core.author = "EZLove 项目组"
    core.keywords = "EZLove, 易挂念, 溪东社区, 老人关怀, 微信小程序, 服务号, AI"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
